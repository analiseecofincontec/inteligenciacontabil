import streamlit as st
import pandas as pd
import openai
import PyPDF2
import io
import os
from docx import Document

# Configurar a API do OpenAI a partir de variáveis de ambiente
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("A variável de ambiente OPENAI_API_KEY não está definida. Configure antes de executar o código.")

# Função para extrair texto de arquivos PDF
def extract_pdf_data(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Função para enviar dados à API do ChatGPT para análise
def analyze_data(data):
    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Você é um analista contábil e financeiro e fará a análise de balancetes e DRE's,\n"
                        "realizará análises horizontais e verticais, cálculos de indicadores contábeis como liquidez, rentabilidade, endividamento e lucratividade.\n"
                        "Além disso, analisa a coerência dos valores de estoque, clientes e fornecedores em relação ao faturamento, e fornece insights precisos sobre a saúde financeira da empresa,\n"
                        "aponte as divergências das contas relevantes que não constam nos relatórios como vendas, custos, água, energia, telefone, aluguel, pró-labore, e outros.\n"
                        "Quando o relatório constar vários períodos sendo meses ou anos, gere automaticamente o gráfico de Receita, Custos, Despesas e Resultado.\n"
                        "Todas as vezes que você calcular um indicador, demonstre a memória de cálculo, isto é muito importante para o usuário.\n"
                        "Pontos de Alerta, quando o Caixa Credor quer dizer que o caixa está estourado, deve-se enfatizar o risco de omissão de receitas e complemente.\n"
                        "Quando constar Adiantamento para Futuro Aumento de Capital, Lançamentos no passivo nessa conta devem ser monitorados e destacados para evitar irregularidades contábeis, enfatize bem os riscos de fazer este lançamento se ele não for de fato real.\n"
                        "Caso o usuário insira informações ou assuntos que não são coerentes com as instruções acima, desconsidere e responda com a mensagem: "
                        "'Sinto muito, mas não fui treinado para responder perguntas deste assunto.'"
                    )
                },
                {"role": "user", "content": f"Por favor, analise os seguintes dados: {data}"}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Erro na análise: {str(e)}"

# Função para salvar o relatório em Word
def save_report_to_word(report_content):
    output = io.BytesIO()
    doc = Document()
    doc.add_heading("Relatório de Análise Contábil e Financeira", level=1)

    report_lines = report_content.split('\n')
    for line in report_lines:
        doc.add_paragraph(line)

    doc.save(output)
    output.seek(0)
    return output

# Aplicação Streamlit
def main():
    st.title("Analisador Contábil e Financeiro")

    st.sidebar.header("Upload de Arquivo")
    uploaded_file = st.sidebar.file_uploader("Envie um arquivo PDF ou Excel", type=["pdf", "xlsx"])

    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            st.info("Processando arquivo PDF...")
            try:
                extracted_text = extract_pdf_data(uploaded_file)
                st.write("Dados extraídos do PDF:")
                st.text_area("Conteúdo extraído", value=extracted_text, height=300)

                st.info("Analisando os dados...")
                report = analyze_data(extracted_text)
                st.write("Relatório gerado:")
                st.text_area("Relatório da Análise", value=report, height=300)

                st.info("Preparando para download...")
                word_data = save_report_to_word(report)
                st.download_button(
                    label="Baixar Relatório",
                    data=word_data,
                    file_name="relatorio_analise.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro ao processar o arquivo PDF: {str(e)}")

        elif uploaded_file.name.endswith(".xlsx"):
            st.info("Processando arquivo Excel...")
            try:
                data = pd.read_excel(uploaded_file)
                st.write("Dados extraídos do Excel:")
                st.dataframe(data)

                st.info("Analisando os dados...")
                report = analyze_data(data.to_string())
                st.write("Relatório gerado:")
                st.text_area("Relatório da Análise", value=report, height=300)

                st.info("Preparando para download...")
                word_data = save_report_to_word(report)
                st.download_button(
                    label="Baixar Relatório",
                    data=word_data,
                    file_name="relatorio_analise.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro ao processar o arquivo Excel: {str(e)}")
    else:
        st.sidebar.info("Por favor, envie um arquivo PDF ou Excel para começar.")

if __name__ == "__main__":
    main()
